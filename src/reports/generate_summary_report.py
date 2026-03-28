#!/usr/bin/env python3
"""
src/reports/generate_summary_report.py

Generates a self-contained Word (.docx) report summarising the Congressional
Stock Analysis pipeline: trade overview, feature EDA, model performance,
ticker analysis, and error analysis.

Usage:
    python src/reports/generate_summary_report.py

Output:
    data/output/summary_report.docx

Prerequisites:
    pip install python-docx
"""

from __future__ import annotations

import os
import sys
import warnings
from io import BytesIO
from datetime import date
from pathlib import Path

warnings.filterwarnings("ignore")

# ── project root ──────────────────────────────────────────────────────────────
PROJECT_ROOT = Path(__file__).resolve().parents[2]
os.chdir(PROJECT_ROOT)
sys.path.insert(0, str(PROJECT_ROOT / "src"))

import numpy as np
import pandas as pd

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import seaborn as sns

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from model.model_realized import PoliticianTradeModel, Config

# ── constants ─────────────────────────────────────────────────────────────────
THRESHOLD_DEFAULT = 0.60
THRESHOLDS        = [0.50, 0.55, 0.60, 0.65, 0.70]
OUTPUT_PATH       = "data/output/summary_report.docx"
MODEL_PATH        = "data/output/xgboost_model_realized.json"

SPY_COLOR   = "#2196F3"
MODEL_COLOR = "#4CAF50"
NAIVE_COLOR = "#FF9800"

sns.set_style("whitegrid")
plt.rcParams.update({"figure.dpi": 100, "font.size": 9})


# ── docx helpers ─────────────────────────────────────────────────────────────
def _fig_to_docx(doc: Document, fig: plt.Figure, width_inches: float = 6.0) -> None:
    buf = BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=110)
    buf.seek(0)
    doc.add_picture(buf, width=Inches(width_inches))
    plt.close(fig)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_text(doc: Document, text: str, italic: bool = False) -> None:
    """Add an explanatory paragraph with consistent body styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    if italic:
        run.italic = True


def _add_table(doc: Document, df: pd.DataFrame) -> None:
    """Render a pandas DataFrame as a Word table."""
    t = doc.add_table(rows=1 + len(df), cols=len(df.columns), style="Table Grid")
    for j, col in enumerate(df.columns):
        cell = t.rows[0].cells[j]
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for row_idx, (_, row) in enumerate(df.iterrows()):
        for j, val in enumerate(row):
            cell = t.rows[row_idx + 1].cells[j]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)


# ── pipeline ─────────────────────────────────────────────────────────────────
def run_pipeline():
    """Reproduce the model train/test split exactly and load the saved model."""
    import xgboost as xgb

    cfg = Config()

    print("Loading enriched parquet …")
    df_full = pd.read_parquet(cfg.data_path)

    print("Calculating sell-pressure features (≈5 min) …")
    df_full = PoliticianTradeModel._calculate_sells_pressure(df_full, "Filed")
    df_full = PoliticianTradeModel._calculate_all_pol_sells_same_ticker(
        df_full, "Filed", window_days=30
    )

    df = df_full[df_full["Transaction"] == "Purchase"].copy()
    df["Traded"] = pd.to_datetime(df["Traded"], errors="coerce")
    max_traded = df["Traded"].max()
    cutoff = (max_traded - pd.DateOffset(months=12)).strftime("%Y-%m-%d")
    print(f"Cutoff date: {cutoff}")

    model = PoliticianTradeModel(cfg, cutoff_date=cutoff, horizon_months=12)

    print("Preprocessing …")
    proc_df = model.preprocess(df)

    train_df, test_df = model.time_split(proc_df)

    train_df = train_df.reset_index(drop=True)
    test_df  = test_df.reset_index(drop=True)

    y_train     = train_df[model.target_binary].reset_index(drop=True)
    y_test      = test_df[model.target_binary].reset_index(drop=True)
    y_test_cont = test_df[model.target_continuous].reset_index(drop=True)

    print("Preparing features …")
    X_train = model.prepare_features(train_df, is_training=True)
    X_test  = model.prepare_features(test_df,  is_training=False)
    X_train = X_train.reset_index(drop=True)
    X_test  = X_test.reset_index(drop=True)

    print("Loading XGBoost model …")
    xgb_clf = xgb.XGBClassifier(enable_categorical=True)
    xgb_clf.load_model(MODEL_PATH)
    y_prob = xgb_clf.predict_proba(X_test)[:, 1]

    return (model, proc_df, train_df, test_df,
            X_train, X_test,
            y_train, y_test, y_test_cont,
            y_prob, df_full, xgb_clf)


# ── Section 1: Trade Overview ─────────────────────────────────────────────────
def section_trade_overview(
    doc: Document,
    df_full: pd.DataFrame,
    proc_df: pd.DataFrame,
) -> None:
    _add_heading(doc, "1. Trade Overview", level=1)

    _add_text(doc,
        "This section summarises the full dataset of congressional stock transactions, "
        "focusing on purchase activity across parties and chambers. The core performance "
        "metric is realized_car_hybrid — the cumulative abnormal return (CAR) calculated "
        "either from the actual buy-to-sell period when the position was closed within "
        "12 months, or from the 12-month fixed window otherwise. A positive value means "
        "the politician's stock outperformed the market benchmark over that horizon."
    )
    doc.add_paragraph()

    purchases = df_full[df_full["Transaction"] == "Purchase"].copy()
    purchases["Traded"] = pd.to_datetime(purchases["Traded"], errors="coerce")

    target    = proc_df["realized_car_hybrid"]
    n_total   = len(df_full)
    n_purchase= len(purchases)
    d_min     = purchases["Traded"].min().strftime("%Y-%m-%d")
    d_max     = purchases["Traded"].max().strftime("%Y-%m-%d")
    pct_pos   = (target > 0).mean()
    mean_car  = target.mean()
    med_car   = target.median()
    std_car   = target.std()

    summary = pd.DataFrame({
        "Metric": [
            "Total records (all transaction types)",
            "Purchases",
            "Date range (traded)",
            "% Positive realized_car_hybrid",
            "Mean realized_car_hybrid",
            "Median realized_car_hybrid",
            "Std dev realized_car_hybrid",
        ],
        "Value": [
            f"{n_total:,}",
            f"{n_purchase:,}",
            f"{d_min} – {d_max}",
            f"{pct_pos:.1%}",
            f"{mean_car:+.4f}",
            f"{med_car:+.4f}",
            f"{std_car:.4f}",
        ],
    })
    _add_table(doc, summary)

    _add_text(doc,
        f"Key finding: the mean realized CAR across all modelled purchases is "
        f"{mean_car:+.4f} ({mean_car*100:+.1f}%), with only {pct_pos:.1%} of trades "
        f"beating the market. The median is {med_car:+.4f}, also below zero. "
        "This confirms that the average congressional stock purchase does not reliably "
        "outperform the benchmark — making the model's task of identifying the positive "
        "tail non-trivial.",
        italic=True,
    )
    doc.add_paragraph()

    # ── 1b: CAR distribution histograms ──────────────────────────────────────
    _add_heading(doc, "CAR Distributions Across Holding Horizons", level=2)

    _add_text(doc,
        "The charts below show the distribution of cumulative abnormal returns (CAR) "
        "for each measurement window. The grey vertical line marks zero (market parity); "
        "the bold coloured line marks the mean for that distribution. Note that the mean "
        "sits below zero across all horizons — indicating that in aggregate, congressional "
        "purchases do not outperform. The right tail (large positive returns) is heavier "
        "than the left, suggesting occasional outsized winners drag the average upward, "
        "but the median remains negative."
    )
    doc.add_paragraph()

    parse     = PoliticianTradeModel._parse_decimal_comma
    car_specs = [
        ("car_filed_to_1m",     "1-month CAR",             proc_df),
        ("car_filed_to_3m",     "3-month CAR",             proc_df),
        ("car_filed_to_6m",     "6-month CAR",             proc_df),
        ("car_filed_to_9m",     "9-month CAR",             proc_df),
        ("car_filed_to_12m",    "12-month CAR",            proc_df),
        ("realized_car_hybrid", "Realized CAR (hybrid)",   proc_df),
    ]
    mean_colors = ["#E67E22", "#9B59B6", "#1ABC9C", "#E74C3C", "#2980B9", "#C0392B"]

    fig, axes = plt.subplots(2, 3, figsize=(13, 7))
    axes = axes.flatten()
    for ax, (col, lbl, src), mc in zip(axes, car_specs, mean_colors):
        if col not in src.columns:
            ax.set_visible(False)
            continue
        data  = parse(src[col].copy()).dropna().clip(-1, 2)
        mu    = data.mean()
        med   = data.median()
        n_pos = (data > 0).mean()

        # histogram
        ax.hist(data, bins=60, color="#5C85D6", alpha=0.65, edgecolor="white", linewidth=0.3)
        # zero reference line (thin, grey)
        ax.axvline(0, color="grey", linewidth=1.0, linestyle="--", alpha=0.6, zorder=3)
        # mean line (thick, coloured, annotated)
        ymax = ax.get_ylim()[1]
        ax.axvline(mu, color=mc, linewidth=2.5, linestyle="-", zorder=4,
                   label=f"Mean = {mu:+.3f}")
        ax.axvline(med, color=mc, linewidth=1.2, linestyle=":", zorder=4,
                   label=f"Median = {med:+.3f}")
        # text annotation inside chart
        ax.text(0.97, 0.95, f"Mean  {mu:+.3f}\nMedian {med:+.3f}\n%>0  {n_pos:.0%}",
                transform=ax.transAxes, fontsize=7, va="top", ha="right",
                bbox=dict(boxstyle="round,pad=0.3", fc="white", alpha=0.8),
                color=mc)
        ax.set_title(lbl, fontsize=9, fontweight="bold")
        ax.set_xlabel("CAR", fontsize=8)
        ax.set_ylabel("Count", fontsize=8)
        ax.tick_params(labelsize=7)

    fig.suptitle("CAR Distributions — Congressional Stock Purchases\n"
                 "(clipped at –1 / +2 for readability; coloured line = mean, dotted = median)",
                 fontsize=10, fontweight="bold")
    fig.tight_layout()
    _fig_to_docx(doc, fig, width_inches=6.5)

    # ── 1c: breakdown by Party / Chamber ─────────────────────────────────────
    _add_heading(doc, "% Positive CAR by Party and Chamber", level=2)

    _add_text(doc,
        "The table below breaks down trade performance by party affiliation and chamber. "
        "Differences between groups reveal whether structural factors (majority party, "
        "access to classified information in Senate committees, etc.) correlate with "
        "better market timing, independent of stock selection skill."
    )
    doc.add_paragraph()

    bd = proc_df[["Party", "Chamber", "realized_car_hybrid"]].copy()
    bd["positive"] = (bd["realized_car_hybrid"] > 0).astype(int)

    for grp_col, label in [("Party", "By Party"), ("Chamber", "By Chamber")]:
        stats = (
            bd.groupby(grp_col)
            .agg(N=("positive", "count"),
                 pct_pos=("positive", "mean"),
                 mean_car=("realized_car_hybrid", "mean"),
                 median_car=("realized_car_hybrid", "median"))
            .reset_index()
        )
        stats["pct_pos"]    = stats["pct_pos"].map("{:.1%}".format)
        stats["mean_car"]   = stats["mean_car"].map("{:+.4f}".format)
        stats["median_car"] = stats["median_car"].map("{:+.4f}".format)
        stats.columns       = [grp_col, "N", "% Positive", "Mean CAR", "Median CAR"]
        doc.add_paragraph(f"{label}:")
        _add_table(doc, stats)
        doc.add_paragraph()

    # ── 1d: quarterly trend ───────────────────────────────────────────────────
    _add_heading(doc, "Trade Volume & % Positive CAR Over Time", level=2)

    _add_text(doc,
        "The chart below tracks quarterly purchase volume (bars, left axis) and the "
        "fraction of those purchases that ended up with a positive realized CAR (line, "
        "right axis). Spikes in the positive-rate line around major market events "
        "(e.g. post-COVID recovery 2020–2021) suggest that congressional timing "
        "may be partly driven by macro conditions rather than stock-specific insight. "
        "Periods of low positive rate often coincide with broad market drawdowns."
    )
    doc.add_paragraph()

    trend = proc_df[["Filed", "realized_car_hybrid"]].copy()
    trend["Filed"]    = pd.to_datetime(trend["Filed"], errors="coerce")
    trend["positive"] = (trend["realized_car_hybrid"] > 0).astype(int)
    trend = trend.dropna(subset=["Filed"])
    trend["quarter"] = trend["Filed"].dt.to_period("Q")

    q = (trend.groupby("quarter")
         .agg(volume=("positive", "count"), pct_pos=("positive", "mean"))
         .reset_index())
    q = q[q["quarter"] >= pd.Period("2015Q1")].copy()
    q["qs"] = q["quarter"].astype(str)

    fig, ax1 = plt.subplots(figsize=(13, 4))
    xs = range(len(q))
    ax1.bar(xs, q["volume"], color="#5C85D6", alpha=0.7, label="Trade Volume")
    ax1.set_ylabel("Trade Volume", color="#5C85D6", fontsize=9)
    ax1.tick_params(axis="y", labelcolor="#5C85D6")

    ax2 = ax1.twinx()
    ax2.plot(xs, q["pct_pos"] * 100, color="#E74C3C", linewidth=2,
             marker="o", markersize=3, label="% Positive")
    ax2.axhline(50, color="#E74C3C", linewidth=0.6, linestyle="--", alpha=0.4)
    ax2.set_ylabel("% Positive CAR", color="#E74C3C", fontsize=9)
    ax2.tick_params(axis="y", labelcolor="#E74C3C")
    ax2.set_ylim(0, 100)

    yr_xs     = [i for i, s in enumerate(q["qs"]) if s.endswith("Q1")]
    yr_labels = [q["qs"].iloc[i][:4] for i in yr_xs]
    ax1.set_xticks(yr_xs)
    ax1.set_xticklabels(yr_labels, rotation=45, fontsize=8)
    ax1.set_title("Quarterly Purchase Volume & % Positive CAR (2015–)",
                  fontsize=10, fontweight="bold")

    lines1, labs1 = ax1.get_legend_handles_labels()
    lines2, labs2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labs1 + labs2, loc="upper left", fontsize=8)
    fig.tight_layout()
    _fig_to_docx(doc, fig, width_inches=6.5)


# ── Section 2: EDA — Features vs Target ──────────────────────────────────────
def section_features_eda(
    doc: Document,
    model: PoliticianTradeModel,
    proc_df: pd.DataFrame,
) -> None:
    _add_heading(doc, "2. Exploratory Data Analysis — Features vs Target", level=1)

    _add_text(doc,
        "This section investigates which measurable factors correlate with positive "
        "realized CAR. The goal is to build intuition about what drives outperformance "
        "before examining the model's learned behaviour in Section 3. We look at "
        "four categories: (A) politician skill signals, (B) committee membership and "
        "lobbying activity, (C) stock-level characteristics, and (D) numerical feature "
        "correlations with the target."
    )
    doc.add_paragraph()

    parse = PoliticianTradeModel._parse_decimal_comma

    # ── 2a: Politician skill quintile analysis ────────────────────────────────
    _add_heading(doc, "2a. Politician Skill Signals", level=2)

    _add_text(doc,
        "The two most direct measures of politician trading skill are: "
        "(1) politician_hit_rate_past — the fraction of that politician's prior purchases "
        "that achieved a positive CAR, and "
        "(2) politician_mean_car_past — the average CAR on their prior purchases. "
        "Both are computed with a look-ahead gap: a past trade's outcome is only counted "
        "once the filing date + 12 months has elapsed, preventing data leakage. "
        "The charts below split trades into quintiles of each signal and show "
        "how mean CAR and the % positive rate change across quintiles."
    )
    doc.add_paragraph()

    skill_feats = [
        ("politician_hit_rate_past",          "Politician Hit Rate (Past)"),
        ("politician_mean_car_past",          "Politician Mean CAR (Past)"),
        ("politician_mean_realized_car_past", "Politician Mean Realized CAR (Past)"),
        ("politician_trades_last_year",       "Trades by Politician (Last Year)"),
    ]

    fig, axes = plt.subplots(1, 4, figsize=(15, 5))
    for ax, (feat, lbl) in zip(axes, skill_feats):
        if feat not in proc_df.columns:
            ax.set_visible(False)
            continue

        tmp = proc_df[[feat, "realized_car_hybrid"]].copy()
        tmp[feat] = parse(tmp[feat].copy())
        tmp = tmp.dropna()
        if len(tmp) < 50:
            ax.set_visible(False)
            continue

        try:
            tmp["q"] = pd.qcut(tmp[feat], q=5, duplicates="drop", labels=False)
        except Exception:
            ax.set_visible(False)
            continue

        q_means  = tmp.groupby("q")["realized_car_hybrid"].mean()
        q_pos    = tmp.groupby("q")["realized_car_hybrid"].apply(lambda s: (s > 0).mean() * 100)
        q_counts = tmp.groupby("q")["realized_car_hybrid"].count()

        ax2 = ax.twinx()
        bars = ax.bar(q_means.index, q_means.values, color="#5C85D6", alpha=0.75,
                      label="Mean CAR")
        ax2.plot(q_pos.index, q_pos.values, "o-", color="#E74C3C",
                 linewidth=2, markersize=5, label="% Positive")
        ax.axhline(0, color="black", linewidth=0.7, linestyle="--", alpha=0.5)

        ax.set_title(lbl, fontsize=8, fontweight="bold")
        ax.set_xlabel("Quintile (1=Low … 5=High)", fontsize=7)
        ax.set_ylabel("Mean CAR", color="#5C85D6", fontsize=7)
        ax2.set_ylabel("% Positive", color="#E74C3C", fontsize=7)
        ax.tick_params(labelsize=7)
        ax2.tick_params(labelsize=7)

        # annotate N per bar
        for i, (bar, n) in enumerate(zip(bars, q_counts)):
            ax.annotate(f"n={n}", xy=(bar.get_x() + bar.get_width()/2, bar.get_height()),
                        xytext=(0, 3), textcoords="offset points", ha="center", fontsize=5)

    fig.suptitle("Quintile Analysis: Politician Skill Metrics vs Realized CAR",
                 fontsize=10, fontweight="bold")
    fig.tight_layout()
    _fig_to_docx(doc, fig, width_inches=6.5)

    _add_text(doc,
        "Interpretation: if the quintile charts show a monotonically increasing pattern "
        "(both mean CAR and % positive rising from Q1 to Q5), that confirms the skill "
        "signal has predictive value. A flat or noisy pattern suggests the signal may be "
        "weak or overshadowed by market conditions during training.",
        italic=True,
    )
    doc.add_paragraph()

    # ── 2b: Committee membership & lobbying ───────────────────────────────────
    _add_heading(doc, "2b. Committee Membership and Lobbying Activity", level=2)

    _add_text(doc,
        "Congress members on committees overseeing specific industries may have "
        "informational advantages when trading in those sectors. The chart below shows "
        "the positive CAR rate for trades made by politicians on each of the nine "
        "committee category groups, split by whether the flag is 1 (on committee) "
        "or 0 (not on committee). The is_committee_chair and is_committee_majority "
        "flags capture additional seniority and political-majority effects. "
        "lobbied_any_90d indicates whether the traded company filed a lobbying "
        "disclosure in the 90 days prior to the trade — a sign of active Washington "
        "engagement that tends to correlate with higher institutional visibility."
    )
    doc.add_paragraph()

    flags = [
        "committee_defense_security", "committee_finance_housing",
        "committee_fiscal_policy",    "committee_energy_environment",
        "committee_health_labor",     "committee_commerce_technology",
        "committee_agriculture",      "committee_infrastructure",
        "committee_oversight",
        "is_committee_chair", "is_committee_majority", "lobbied_any_90d",
    ]
    rows = []
    for flag in flags:
        if flag not in proc_df.columns:
            continue
        col = pd.to_numeric(proc_df[flag], errors="coerce").fillna(0)
        g1  = proc_df.loc[col == 1, "realized_car_hybrid"]
        g0  = proc_df.loc[col == 0, "realized_car_hybrid"]
        rows.append({
            "flag":  flag.replace("committee_", "").replace("_", " ").title(),
            "pct1":  (g1 > 0).mean() * 100 if len(g1) > 0 else 0.0,
            "pct0":  (g0 > 0).mean() * 100 if len(g0) > 0 else 0.0,
            "mean1": g1.mean() if len(g1) > 0 else 0.0,
            "mean0": g0.mean() if len(g0) > 0 else 0.0,
            "n1":    len(g1),
        })
    flag_df = pd.DataFrame(rows)

    fig, axes = plt.subplots(1, 2, figsize=(13, 5))

    # Left: % positive comparison
    xs = range(len(flag_df))
    w  = 0.38
    axes[0].bar([i - w/2 for i in xs], flag_df["pct1"], w,
                label="Flag = 1 (on committee)", color="#E74C3C", alpha=0.8)
    axes[0].bar([i + w/2 for i in xs], flag_df["pct0"], w,
                label="Flag = 0", color="#5C85D6", alpha=0.8)
    axes[0].axhline(50, color="gray", linewidth=0.7, linestyle="--", alpha=0.5)
    axes[0].set_xticks(list(xs))
    axes[0].set_xticklabels(flag_df["flag"], rotation=40, ha="right", fontsize=7)
    axes[0].set_ylabel("% Positive CAR", fontsize=9)
    axes[0].set_ylim(0, 85)
    axes[0].set_title("% Positive CAR: Flag = 1 vs 0", fontsize=9, fontweight="bold")
    axes[0].legend(fontsize=8)

    # Right: mean CAR comparison
    axes[1].bar([i - w/2 for i in xs], flag_df["mean1"], w,
                label="Flag = 1", color="#E74C3C", alpha=0.8)
    axes[1].bar([i + w/2 for i in xs], flag_df["mean0"], w,
                label="Flag = 0", color="#5C85D6", alpha=0.8)
    axes[1].axhline(0, color="black", linewidth=0.8, linestyle="--", alpha=0.6)
    axes[1].set_xticks(list(xs))
    axes[1].set_xticklabels(flag_df["flag"], rotation=40, ha="right", fontsize=7)
    axes[1].set_ylabel("Mean Realized CAR", fontsize=9)
    axes[1].set_title("Mean Realized CAR: Flag = 1 vs 0", fontsize=9, fontweight="bold")
    axes[1].legend(fontsize=8)

    fig.suptitle("Committee Membership & Lobbying Activity vs Trading Performance",
                 fontsize=10, fontweight="bold")
    fig.tight_layout()
    _fig_to_docx(doc, fig, width_inches=6.5)

    # Lift table
    lift_tbl = flag_df[["flag", "pct1", "pct0", "mean1", "mean0", "n1"]].copy()
    lift_tbl["lift_%pos"] = (lift_tbl["pct1"] - lift_tbl["pct0"]).map("{:+.1f}pp".format)
    lift_tbl["pct1"]  = lift_tbl["pct1"].map("{:.1f}%".format)
    lift_tbl["pct0"]  = lift_tbl["pct0"].map("{:.1f}%".format)
    lift_tbl["mean1"] = lift_tbl["mean1"].map("{:+.4f}".format)
    lift_tbl["mean0"] = lift_tbl["mean0"].map("{:+.4f}".format)
    lift_tbl.columns  = ["Flag", "%Pos (Flag=1)", "%Pos (Flag=0)",
                          "Mean CAR (1)", "Mean CAR (0)", "N (Flag=1)", "Lift (%Pos)"]
    _add_table(doc, lift_tbl)

    _add_text(doc,
        "Positive lift values indicate that being on a given committee category is "
        "associated with higher positive-CAR rates. A large lift on lobbied_any_90d "
        "suggests that actively lobbied companies may be better-known large-caps "
        "rather than a direct insider-information signal.",
        italic=True,
    )
    doc.add_paragraph()

    # ── 2c: Stock-level signals ───────────────────────────────────────────────
    _add_heading(doc, "2c. Stock-Level Characteristics", level=2)

    _add_text(doc,
        "Beyond politician-specific signals, the stock's own momentum and volatility "
        "at the time of purchase may predict future performance. High momentum stocks "
        "could indicate trend-following behaviour; high volatility introduces noise. "
        "Beta captures systematic market exposure — low-beta defensive stocks may "
        "behave differently under different market regimes."
    )
    doc.add_paragraph()

    stock_feats = [
        ("stock_momentum_90d",  "Stock Momentum (90d)"),
        ("stock_volatility_30d","Stock Volatility (30d)"),
        ("beta",                "Beta vs Market"),
        ("max_committee_rank",  "Max Committee Rank"),
    ]

    fig, axes = plt.subplots(1, 4, figsize=(15, 5))
    for ax, (feat, lbl) in zip(axes, stock_feats):
        if feat not in proc_df.columns:
            ax.set_visible(False)
            continue

        tmp = proc_df[[feat, "realized_car_hybrid"]].copy()
        tmp[feat] = parse(tmp[feat].copy())
        tmp = tmp.dropna()
        if len(tmp) < 50:
            ax.set_visible(False)
            continue

        try:
            tmp["q"] = pd.qcut(tmp[feat], q=5, duplicates="drop", labels=False)
        except Exception:
            ax.set_visible(False)
            continue

        q_means = tmp.groupby("q")["realized_car_hybrid"].mean()
        q_pos   = tmp.groupby("q")["realized_car_hybrid"].apply(
            lambda s: (s > 0).mean() * 100
        )

        ax2 = ax.twinx()
        ax.bar(q_means.index, q_means.values, color="#2ECC71", alpha=0.75)
        ax2.plot(q_pos.index, q_pos.values, "o-", color="#E74C3C",
                 linewidth=2, markersize=5)
        ax.axhline(0, color="black", linewidth=0.7, linestyle="--", alpha=0.5)
        ax.set_title(lbl, fontsize=8, fontweight="bold")
        ax.set_xlabel("Quintile (1=Low … 5=High)", fontsize=7)
        ax.set_ylabel("Mean CAR", color="#2ECC71", fontsize=7)
        ax2.set_ylabel("% Positive", color="#E74C3C", fontsize=7)
        ax.tick_params(labelsize=7)
        ax2.tick_params(labelsize=7)

    fig.suptitle("Quintile Analysis: Stock Characteristics vs Realized CAR",
                 fontsize=10, fontweight="bold")
    fig.tight_layout()
    _fig_to_docx(doc, fig, width_inches=6.5)

    _add_text(doc,
        "Note: for momentum, a positive slope from Q1 to Q5 would indicate that "
        "politicians are buying into rising stocks that continue to outperform "
        "(momentum effect). A U-shaped or inverted pattern would suggest mean "
        "reversion dynamics. Volatility quintiles help assess whether the model "
        "should favour lower-risk, steadier outperformers.",
        italic=True,
    )
    doc.add_paragraph()

    # ── 2d: Numerical feature correlation with target ──────────────────────────
    _add_heading(doc, "2d. Feature Correlations with Realized CAR", level=2)

    _add_text(doc,
        "The table below shows the Pearson and Spearman correlations between each "
        "numerical feature and the realized_car_hybrid target, computed on the full "
        "modelled dataset (before train/test split). Pearson captures linear relationships; "
        "Spearman captures monotonic ones. Features with high absolute Spearman "
        "correlation are the most consistently predictive across the distribution."
    )
    doc.add_paragraph()

    num_feats = [
        "politician_hit_rate_past", "politician_mean_car_past",
        "politician_mean_realized_car_past", "politician_trades_last_year",
        "stock_momentum_90d", "stock_momentum_30d", "stock_volatility_30d",
        "beta", "max_committee_rank", "log_trade_size",
        "car_traded_to_filed", "n_committees", "ticker_prior_buys",
        "all_pol_sells_same_ticker_30d",
    ]

    corr_rows = []
    for feat in num_feats:
        if feat not in proc_df.columns:
            continue
        x = parse(proc_df[feat].copy()).dropna()
        y = proc_df.loc[x.index, "realized_car_hybrid"].dropna()
        common = x.index.intersection(y.index)
        if len(common) < 30:
            continue
        xc, yc = x.loc[common], y.loc[common]
        pearson  = xc.corr(yc)
        spearman = xc.rank().corr(yc.rank())
        corr_rows.append({
            "Feature": feat,
            "Pearson r":   f"{pearson:+.4f}",
            "Spearman ρ":  f"{spearman:+.4f}",
            "|Spearman|":  abs(spearman),
        })

    if corr_rows:
        corr_df = (pd.DataFrame(corr_rows)
                   .sort_values("|Spearman|", ascending=False)
                   .drop(columns=["|Spearman|"]))
        _add_table(doc, corr_df)

    _add_text(doc,
        "Features at the top of this table provide the strongest monotonic signal "
        "for predicting above-market returns. Even modest correlations (~0.05–0.15) "
        "can be economically significant when combined in an ensemble model, especially "
        "if they are orthogonal (i.e. each captures a different aspect of the trade).",
        italic=True,
    )


# ── Section 3: Model Performance ─────────────────────────────────────────────
def section_model_performance(
    doc: Document,
    model: PoliticianTradeModel,
    proc_df: pd.DataFrame,
    test_df: pd.DataFrame,
    y_test: pd.Series,
    y_test_cont: pd.Series,
    y_prob: np.ndarray,
    xgb_clf,
) -> None:
    from sklearn.metrics import (precision_score, recall_score,
                                  f1_score, confusion_matrix)

    _add_heading(doc, "3. Model Performance & Feature Importances", level=1)

    _add_text(doc,
        "The model is an XGBoost classifier trained on an 80/20 time-based split "
        "(never random, always chronological) to prevent look-ahead bias. The target "
        "is a binary indicator: does the realized_car_hybrid exceed 0? "
        "Model inputs are 31 features spanning politician skill, stock characteristics, "
        "committee membership, and lobbying signals. "
        "This section covers (A) which features the model learned to rely on most, "
        "followed by (B) threshold-level precision/recall trade-offs, "
        "(C) the confusion matrix, and (D) a dollar-value simulation on the test set."
    )
    doc.add_paragraph()

    # ── 3a: feature importance (moved from Section 2) ─────────────────────────
    _add_heading(doc, "3a. XGBoost Feature Importances", level=2)

    _add_text(doc,
        "Feature importance (measured by average gain across all splits using each "
        "feature) shows which signals the model weighted most heavily. Features at "
        "the top explain a disproportionate share of the model's predictive power. "
        "It is notable whether politician-specific skill signals (hit rate, mean CAR) "
        "dominate over market signals (momentum, volatility) — this has implications "
        "for the source of any alpha in the model."
    )
    doc.add_paragraph()

    fi    = pd.Series(xgb_clf.feature_importances_, index=model.feature_names).sort_values(ascending=False)
    top20 = fi.head(20).sort_values()

    fig, ax = plt.subplots(figsize=(8, 6))
    colors = ["#E74C3C" if v >= top20.median() else "#5C85D6" for v in top20.values]
    top20.plot(kind="barh", ax=ax, color=colors, edgecolor="white", linewidth=0.4)
    ax.set_title("Top 20 XGBoost Feature Importances — Realized CAR (gain)",
                 fontsize=10, fontweight="bold")
    ax.set_xlabel("Importance (gain)", fontsize=9)
    ax.tick_params(labelsize=8)
    fig.tight_layout()
    _fig_to_docx(doc, fig, width_inches=6.2)

    # importance table (top 10)
    fi_tbl = fi.head(10).reset_index()
    fi_tbl.columns = ["Feature", "Importance (gain)"]
    fi_tbl["Importance (gain)"] = fi_tbl["Importance (gain)"].map("{:.6f}".format)
    _add_table(doc, fi_tbl)

    _add_text(doc,
        "Red bars mark features above the median importance threshold. "
        "A model heavily dominated by one or two features may be overfitting to "
        "that signal; a more distributed importance profile suggests the model is "
        "combining complementary signals.",
        italic=True,
    )
    doc.add_paragraph()

    # ── 3b: threshold table ───────────────────────────────────────────────────
    _add_heading(doc, "3b. Threshold Performance Table", level=2)

    _add_text(doc,
        "Adjusting the classification threshold changes the precision/recall trade-off. "
        "A higher threshold means the model only signals a trade when it is highly "
        "confident, reducing the number of predictions but increasing precision. "
        "The 'Mean CAR (Pred +)' column shows the average return of the trades flagged "
        "as positive at each threshold — this is the most economically relevant metric."
    )
    doc.add_paragraph()

    split_dt = proc_df.sort_values("Filed")[model.date_column].iloc[int(len(proc_df) * 0.8)]
    _add_text(doc,
        f"Train/test split date: {pd.Timestamp(split_dt).date()}  "
        f"| Test set size: {len(test_df):,} trades",
        italic=True,
    )
    doc.add_paragraph()

    perf_rows = []
    for thr in THRESHOLDS:
        y_pred  = (y_prob >= thr).astype(int)
        mask_pp = y_pred == 1
        perf_rows.append({
            "Threshold":         f"{thr:.2f}",
            "Precision":         f"{precision_score(y_test, y_pred, zero_division=0):.3f}",
            "Recall":            f"{recall_score(y_test, y_pred, zero_division=0):.3f}",
            "F1":                f"{f1_score(y_test, y_pred, zero_division=0):.3f}",
            "N Predictions":     f"{int(y_pred.sum()):,}",
            "Mean CAR (Pred +)": (f"{y_test_cont[mask_pp].mean():+.4f}"
                                  if mask_pp.sum() > 0 else "—"),
        })
    _add_table(doc, pd.DataFrame(perf_rows))

    _add_text(doc,
        "Precision is the fraction of predicted positives that actually beat the market. "
        "At threshold 0.60 the model offers a balance between coverage (recall) and "
        "accuracy (precision). The mean CAR of predicted positives is the key figure: "
        "any value meaningfully above the test-set baseline indicates genuine lift.",
        italic=True,
    )
    doc.add_paragraph()

    # ── 3c: confusion matrix ──────────────────────────────────────────────────
    _add_heading(doc, f"3c. Confusion Matrix (threshold = {THRESHOLD_DEFAULT})", level=2)

    _add_text(doc,
        f"At the default threshold of {THRESHOLD_DEFAULT}, the confusion matrix "
        "below shows the counts of true positives, false positives, true negatives, "
        "and false negatives on the test set. A high false-positive count is costly "
        "in practical terms (investing in trades that underperform); false negatives "
        "represent missed opportunities."
    )
    doc.add_paragraph()

    y_pred60 = (y_prob >= THRESHOLD_DEFAULT).astype(int)
    cm = confusion_matrix(y_test, y_pred60)

    fig, ax = plt.subplots(figsize=(4.5, 3.5))
    sns.heatmap(cm, annot=True, fmt="d", cmap="Blues", cbar=False,
                xticklabels=["Pred 0", "Pred 1"],
                yticklabels=["Actual 0", "Actual 1"], ax=ax,
                annot_kws={"size": 12})
    ax.set_title(f"Confusion Matrix @ threshold = {THRESHOLD_DEFAULT}",
                 fontsize=10, fontweight="bold")
    fig.tight_layout()
    _fig_to_docx(doc, fig, width_inches=3.8)

    # ── 3d: $100-per-trade simulation ─────────────────────────────────────────
    _add_heading(doc, "3d. $100-per-Trade Cumulative Return Simulation", level=2)

    _add_text(doc,
        "To put model performance in practical context, the simulation below assumes "
        "an investor places $100 into each trade flagged as positive by the model "
        f"(threshold = {THRESHOLD_DEFAULT}), or $100 into every purchase for the naive "
        "baseline. The y-axis shows cumulative total receipts (principal + return) "
        "as trades are entered in chronological order through the test set. The SPY "
        "benchmark invests $100 on each of the same trade dates and measures the "
        "12-month forward return on the S&P 500 ETF, providing a like-for-like "
        "comparison against a passive index strategy."
    )
    doc.add_paragraph()

    sim = test_df[["Filed", "realized_car_hybrid"]].copy()
    sim["Filed"]   = pd.to_datetime(sim["Filed"], errors="coerce")
    sim["y_prob"]  = y_prob
    sim = sim.dropna(subset=["Filed", "realized_car_hybrid"])
    sim = sim.sort_values("Filed").reset_index(drop=True)

    mask_model     = sim["y_prob"] >= THRESHOLD_DEFAULT
    sim["ret_model"]= np.where(mask_model, 100.0 * (1 + sim["realized_car_hybrid"]), 0.0)
    sim["ret_naive"]= 100.0 * (1 + sim["realized_car_hybrid"])
    sim["cum_model"]= sim["ret_model"].cumsum()
    sim["cum_naive"]= sim["ret_naive"].cumsum()

    has_spy = False
    try:
        import yfinance as yf
        spy_raw = yf.download(
            "SPY",
            start=(sim["Filed"].min() - pd.Timedelta(days=5)).strftime("%Y-%m-%d"),
            end=(sim["Filed"].max() + pd.Timedelta(days=400)).strftime("%Y-%m-%d"),
            progress=False,
        )
        spy_close = (spy_raw[("Close", "SPY")]
                     if isinstance(spy_raw.columns, pd.MultiIndex)
                     else spy_raw["Close"]).dropna()

        def _spy_fwd(d):
            try:
                after  = spy_close.index[spy_close.index >= d]
                before = spy_close.index[spy_close.index <= d + pd.Timedelta(days=365)]
                if len(after) == 0 or len(before) == 0:
                    return np.nan
                return float(spy_close.loc[before[-1]] - spy_close.loc[after[0]]) / float(spy_close.loc[after[0]])
            except Exception:
                return np.nan

        sim["spy_ret"]  = sim["Filed"].apply(_spy_fwd)
        sim["ret_spy"]  = 100.0 * (1 + sim["spy_ret"].fillna(0))
        sim["cum_spy"]  = sim["ret_spy"].cumsum()
        has_spy = True
    except Exception as e:
        print(f"  SPY data unavailable: {e}")

    fig, ax = plt.subplots(figsize=(13, 5))
    ax.plot(sim["Filed"], sim["cum_naive"], color=NAIVE_COLOR, linewidth=1.5,
            label="Naive baseline — all purchases")
    ax.plot(sim["Filed"], sim["cum_model"], color=MODEL_COLOR, linewidth=2.2,
            label=f"Model — threshold = {THRESHOLD_DEFAULT}")
    if has_spy:
        ax.plot(sim["Filed"], sim["cum_spy"], color=SPY_COLOR, linewidth=1.5,
                linestyle="--", label="SPY benchmark (same dates, 12m window)")

    for val, color in [(sim["cum_naive"].iloc[-1], NAIVE_COLOR),
                       (sim["cum_model"].iloc[-1], MODEL_COLOR)]:
        ax.annotate(f"${val:,.0f}", xy=(sim["Filed"].iloc[-1], val),
                    xytext=(-5, 4), textcoords="offset points",
                    ha="right", fontsize=8, color=color, fontweight="bold")
    if has_spy:
        val = sim["cum_spy"].iloc[-1]
        ax.annotate(f"${val:,.0f}", xy=(sim["Filed"].iloc[-1], val),
                    xytext=(-5, -10), textcoords="offset points",
                    ha="right", fontsize=8, color=SPY_COLOR, fontweight="bold")

    ax.set_xlabel("Trade Filing Date", fontsize=9)
    ax.set_ylabel("Cumulative Portfolio Value ($)", fontsize=9)
    ax.set_title("$100-per-Trade Cumulative Return — Test Set (chronological)",
                 fontsize=10, fontweight="bold")
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"${x:,.0f}"))
    ax.legend(fontsize=8)
    fig.tight_layout()
    _fig_to_docx(doc, fig, width_inches=6.5)

    n_model  = int(mask_model.sum())
    n_naive  = len(sim)
    invested = n_model * 100
    fin_model = sim["cum_model"].iloc[-1]
    fin_naive = sim["cum_naive"].iloc[-1]
    _add_text(doc,
        f"The model selected {n_model:,} trades ({n_model/n_naive:.1%} of the test set), "
        f"investing ${invested:,} for a final cumulative value of ${fin_model:,.0f} "
        f"(return: {(fin_model - invested) / invested * 100:+.1f}%). "
        f"The naive baseline invested ${n_naive * 100:,} across all {n_naive:,} "
        f"test purchases and accumulated ${fin_naive:,.0f} "
        f"(return: {(fin_naive - n_naive*100) / (n_naive*100) * 100:+.1f}%). "
        "Interpret with caution: this simulation assumes no transaction costs, "
        "full execution at filing date, and no position-size constraints.",
        italic=True,
    )


# ── Section 4: Predicted Positive Ticker Analysis ─────────────────────────────
def section_ticker_analysis(
    doc: Document,
    test_df: pd.DataFrame,
    y_test: pd.Series,
    y_test_cont: pd.Series,
    y_prob: np.ndarray,
) -> None:
    _add_heading(doc, "4. Predicted Positive Ticker Analysis", level=1)

    _add_text(doc,
        "This section examines which stocks the model most frequently identifies as "
        "likely outperformers, and whether those predictions are actually correct. "
        "Concentration in a small number of tickers could indicate the model has "
        "learned issuer-specific patterns (e.g. a politician who consistently trades "
        "one stock well) rather than generalisable signals. Sector-level analysis "
        "reveals whether the model overweights certain industries."
    )
    doc.add_paragraph()

    pp_idx   = np.where(y_prob >= THRESHOLD_DEFAULT)[0]
    pred_pos = test_df.iloc[pp_idx].copy()
    pred_pos["y_prob"]    = y_prob[pp_idx]
    pred_pos["y_actual"]  = y_test.iloc[pp_idx].values
    pred_pos["car"]       = y_test_cont.iloc[pp_idx].values
    pred_pos["pred_type"] = np.where(pred_pos["y_actual"] == 1, "TP", "FP")

    # ── 4a: top-20 tickers ────────────────────────────────────────────────────
    _add_heading(doc, "Top 20 Tickers Most Frequently Predicted Positive", level=2)

    _add_text(doc,
        "Each bar represents the count of test-set trades where the model predicted "
        "positive for that ticker. The orange line (right axis) shows the actual "
        "mean realized CAR for those predicted trades. Tickers with many predictions "
        "but a negative mean CAR are systematically false-positived by the model."
    )
    doc.add_paragraph()

    tc = pred_pos["Ticker"].value_counts().head(20)
    mc = pred_pos.groupby("Ticker")["car"].mean()

    fig, ax1 = plt.subplots(figsize=(11, 5))
    xs = range(len(tc))
    ax1.bar(xs, tc.values, color="#5C85D6", alpha=0.8, label="# Predictions")
    ax1.set_ylabel("Number of Predictions", color="#5C85D6", fontsize=9)
    ax1.tick_params(axis="y", labelcolor="#5C85D6")

    ax2 = ax1.twinx()
    mean_cars = [mc.get(t, 0.0) for t in tc.index]
    ax2.plot(xs, mean_cars, "o-", color="#E74C3C", linewidth=2, markersize=5,
             label="Mean Realized CAR")
    ax2.axhline(0, color="gray", linewidth=0.6, linestyle="--", alpha=0.5)
    ax2.set_ylabel("Mean Realized CAR", color="#E74C3C", fontsize=9)
    ax2.tick_params(axis="y", labelcolor="#E74C3C")

    ax1.set_xticks(list(xs))
    ax1.set_xticklabels(tc.index, rotation=45, ha="right", fontsize=8)
    ax1.set_title("Top 20 Tickers Predicted Positive (threshold = 0.60)",
                  fontsize=10, fontweight="bold")
    lines1, labs1 = ax1.get_legend_handles_labels()
    lines2, labs2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labs1 + labs2, loc="upper right", fontsize=8)
    fig.tight_layout()
    _fig_to_docx(doc, fig, width_inches=6.5)

    # ── 4b: sector breakdown ──────────────────────────────────────────────────
    _add_heading(doc, "Sector Distribution: Predicted Positives vs All Test Trades", level=2)

    _add_text(doc,
        "If the model's predicted positives mirror the sector distribution of all test "
        "trades, it is sector-agnostic. Deviations show sector tilt — whether by design "
        "(committee-linked sectors genuinely outperform) or by overfitting to historical "
        "sector cycles within the training window."
    )
    doc.add_paragraph()

    all_sec  = test_df["Ticker_Sector"].fillna("Unknown").value_counts()
    pred_sec = pred_pos["Ticker_Sector"].fillna("Unknown").value_counts()
    sectors  = all_sec.index.tolist()
    all_pct  = (all_sec  / all_sec.sum()  * 100).reindex(sectors, fill_value=0)
    pred_pct = (pred_sec / pred_sec.sum() * 100).reindex(sectors, fill_value=0)

    fig, ax = plt.subplots(figsize=(11, 4.5))
    xs = range(len(sectors))
    w  = 0.38
    ax.bar([i - w/2 for i in xs], all_pct.values,  w,
           label="All Test Trades", color="#5C85D6", alpha=0.8)
    ax.bar([i + w/2 for i in xs], pred_pct.values, w,
           label="Predicted Positives", color="#E74C3C", alpha=0.8)
    ax.set_xticks(list(xs))
    ax.set_xticklabels(sectors, rotation=40, ha="right", fontsize=8)
    ax.set_ylabel("% of Group", fontsize=9)
    ax.set_title("Sector Distribution: Predicted Positives vs All Test Trades",
                 fontsize=10, fontweight="bold")
    ax.legend(fontsize=8)
    fig.tight_layout()
    _fig_to_docx(doc, fig, width_inches=6.5)

    # ── 4c: scatter ───────────────────────────────────────────────────────────
    _add_heading(doc, "Predicted Probability vs Actual CAR (TP / FP)", level=2)

    _add_text(doc,
        "Each point is a trade predicted positive by the model. Green points (TP) "
        "are trades that actually beat the market; red points (FP) did not. "
        "A well-calibrated model should show green points more densely at high "
        "probabilities and more red points at lower probabilities just above the "
        f"{THRESHOLD_DEFAULT} cut-off. The horizontal dashed line marks CAR = 0."
    )
    doc.add_paragraph()

    tp_m = pred_pos["pred_type"] == "TP"
    fp_m = pred_pos["pred_type"] == "FP"

    fig, ax = plt.subplots(figsize=(7, 5))
    ax.scatter(pred_pos.loc[fp_m, "y_prob"],
               pred_pos.loc[fp_m, "car"].clip(-1, 2),
               alpha=0.35, color="#E74C3C", s=12,
               label=f"False Positive (N={fp_m.sum()})")
    ax.scatter(pred_pos.loc[tp_m, "y_prob"],
               pred_pos.loc[tp_m, "car"].clip(-1, 2),
               alpha=0.35, color="#4CAF50", s=12,
               label=f"True Positive (N={tp_m.sum()})")
    ax.axhline(0, color="black", linewidth=0.8, linestyle="--", alpha=0.6)
    ax.set_xlabel("Predicted Probability", fontsize=9)
    ax.set_ylabel("Actual Realized CAR (clipped −1 to +2)", fontsize=9)
    ax.set_title("Predicted Probability vs Actual CAR — Predicted Positives",
                 fontsize=10, fontweight="bold")
    ax.legend(fontsize=8)
    fig.tight_layout()
    _fig_to_docx(doc, fig, width_inches=5.5)


# ── Section 5: Error Analysis ─────────────────────────────────────────────────
def section_error_analysis(
    doc: Document,
    model: PoliticianTradeModel,
    test_df: pd.DataFrame,
    y_test: pd.Series,
    y_test_cont: pd.Series,
    y_prob: np.ndarray,
) -> None:
    _add_heading(doc, "5. Error Analysis", level=1)

    _add_text(doc,
        "Understanding when and why the model makes mistakes is critical for "
        "improving it. This section compares the average feature values for "
        "true positives (TP), false positives (FP), true negatives (TN), and "
        "false negatives (FN). Features where TP mean >> FP mean are those where "
        "the model partially uses the right signal but may need a stronger threshold. "
        "Features where FN mean >> TN mean are the ones the model is failing to "
        "exploit for the missed positive trades."
    )
    doc.add_paragraph()

    y_pred = (y_prob >= THRESHOLD_DEFAULT).astype(int)
    y_true = y_test.astype(int).values

    tp_mask = (y_pred == 1) & (y_true == 1)
    fp_mask = (y_pred == 1) & (y_true == 0)
    tn_mask = (y_pred == 0) & (y_true == 0)
    fn_mask = (y_pred == 0) & (y_true == 1)

    parse    = PoliticianTradeModel._parse_decimal_comma
    num_cols = [c for c in model.numerical_features if c in test_df.columns]

    def _profile(mask):
        if mask.sum() == 0:
            return pd.Series(np.nan, index=num_cols)
        return test_df.iloc[mask][num_cols].apply(lambda s: parse(s.copy()).mean())

    # ── 5a: FP vs TP ──────────────────────────────────────────────────────────
    _add_heading(doc, "False Positive vs True Positive — Feature Means", level=2)

    _add_text(doc,
        f"N True Positives: {tp_mask.sum():,}   |   N False Positives: {fp_mask.sum():,}  "
        f"(threshold = {THRESHOLD_DEFAULT}). "
        "A positive 'TP − FP' difference means true positives score higher on that "
        "feature. Negative differences reveal features where false positives are "
        "unexpectedly high — potential sources of model confusion."
    )
    doc.add_paragraph()

    tp_prof = _profile(tp_mask)
    fp_prof = _profile(fp_mask)
    fp_tp = pd.DataFrame({"TP Mean": tp_prof.round(4), "FP Mean": fp_prof.round(4)})
    fp_tp["TP − FP"] = (fp_tp["TP Mean"] - fp_tp["FP Mean"]).round(4)
    fp_tp = fp_tp.reset_index().rename(columns={"index": "Feature"})
    _add_table(doc, fp_tp)
    doc.add_paragraph()

    # ── 5b: FN vs TN ──────────────────────────────────────────────────────────
    _add_heading(doc, "False Negative vs True Negative — Feature Means", level=2)

    _add_text(doc,
        f"N False Negatives: {fn_mask.sum():,}   |   N True Negatives: {tn_mask.sum():,}. "
        "False negatives are trades the model missed — trades that actually beat the "
        "market but were not flagged. A positive 'FN − TN' difference on a feature "
        "means that missed trades scored higher on that signal, suggesting the model "
        "may not be weighting it sufficiently for borderline cases."
    )
    doc.add_paragraph()

    fn_prof = _profile(fn_mask)
    tn_prof = _profile(tn_mask)
    fn_tn = pd.DataFrame({"FN Mean": fn_prof.round(4), "TN Mean": tn_prof.round(4)})
    fn_tn["FN − TN"] = (fn_tn["FN Mean"] - fn_tn["TN Mean"]).round(4)
    fn_tn = fn_tn.reset_index().rename(columns={"index": "Feature"})
    _add_table(doc, fn_tn)
    doc.add_paragraph()

    # ── 5c: calibration ───────────────────────────────────────────────────────
    _add_heading(doc, "Calibration Plot", level=2)

    _add_text(doc,
        "A well-calibrated model's predicted probabilities should match the observed "
        "positive rate. The diagonal dashed line represents perfect calibration. "
        "If the curve bows above the diagonal the model is under-confident (predicting "
        "lower probabilities than the true rate); below the diagonal indicates "
        "over-confidence. Significant miscalibration at high probabilities is most "
        "impactful because it affects the trades that would actually be selected."
    )
    doc.add_paragraph()

    bins      = np.linspace(0, 1, 11)
    mean_probs, actual_rates = [], []
    for lo, hi in zip(bins[:-1], bins[1:]):
        mask = (y_prob >= lo) & (y_prob < hi)
        if mask.sum() >= 5:
            mean_probs.append(float(y_prob[mask].mean()))
            actual_rates.append(float(y_true[mask].mean()))

    fig, ax = plt.subplots(figsize=(5, 4.5))
    ax.plot([0, 1], [0, 1], "k--", linewidth=1, label="Perfect calibration")
    ax.plot(mean_probs, actual_rates, "o-", color="#E74C3C",
            linewidth=2, markersize=6, label="XGBoost")
    ax.set_xlim(0, 1); ax.set_ylim(0, 1)
    ax.set_xlabel("Mean Predicted Probability", fontsize=9)
    ax.set_ylabel("Actual Positive Rate", fontsize=9)
    ax.set_title("Calibration Plot", fontsize=10, fontweight="bold")
    ax.legend(fontsize=8)
    fig.tight_layout()
    _fig_to_docx(doc, fig, width_inches=4.5)

    # ── 5d: CAR density ───────────────────────────────────────────────────────
    _add_heading(doc, "CAR Distribution by Prediction Type (TP / TN / FP / FN)", level=2)

    _add_text(doc,
        "The overlaid density histograms below show the realized CAR distribution "
        "separately for each of the four prediction outcomes. Ideally, TP trades "
        "should have a distribution shifted strongly to the right; FP trades should "
        "cluster near zero or negative territory, showing the model is picking up "
        "borderline cases. A wide FN distribution that overlaps with TP suggests "
        "the model's false negatives are not obviously different from true positives "
        "— making them genuinely hard to predict."
    )
    doc.add_paragraph()

    groups = [
        ("True Positive",  tp_mask, "#4CAF50"),
        ("True Negative",  tn_mask, "#2196F3"),
        ("False Positive", fp_mask, "#E74C3C"),
        ("False Negative", fn_mask, "#FF9800"),
    ]
    fig, ax = plt.subplots(figsize=(9, 4.5))
    for name, mask, color in groups:
        if mask.sum() < 5:
            continue
        data = y_test_cont.values[mask].clip(-1, 2)
        ax.hist(data, bins=40, alpha=0.45, color=color,
                label=f"{name} (N={mask.sum()})", density=True,
                histtype="stepfilled", edgecolor=color, linewidth=0.5)
    ax.axvline(0, color="black", linewidth=1.0, linestyle="--")
    ax.set_xlabel("Realized CAR Hybrid (clipped −1 to +2)", fontsize=9)
    ax.set_ylabel("Density", fontsize=9)
    ax.set_title("CAR Distribution by Prediction Type", fontsize=10, fontweight="bold")
    ax.legend(fontsize=8)
    fig.tight_layout()
    _fig_to_docx(doc, fig, width_inches=6.2)


# ── main ─────────────────────────────────────────────────────────────────────
def main() -> None:
    print("=" * 70)
    print("  CONGRESSIONAL STOCK ANALYSIS — SUMMARY REPORT GENERATOR")
    print("=" * 70)

    (model, proc_df, train_df, test_df,
     X_train, X_test,
     y_train, y_test, y_test_cont,
     y_prob, df_full, xgb_clf) = run_pipeline()

    print("\nBuilding Word document …")
    doc = Document()

    # ── title page ────────────────────────────────────────────────────────────
    title = doc.add_heading("Congressional Stock Analysis — Summary Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in [
        f"Generated: {date.today().strftime('%B %d, %Y')}",
        f"Dataset: {len(proc_df):,} modelled purchases  |  "
        f"Train: {len(train_df):,}  |  Test: {len(test_df):,}",
        "Model: XGBoost classifier  |  Target: realized_car_hybrid > 0",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    _add_text(doc,
        "This report summarises the end-to-end Congressional Stock Analysis pipeline. "
        "It examines whether U.S. Congress members' equity purchases systematically "
        "outperform the market, identifies which features predict outperformance, "
        "and evaluates an XGBoost classifier trained to flag the most promising trades. "
        "The analysis covers five sections: (1) a descriptive overview of all trades, "
        "(2) exploratory analysis of features relative to the return target, "
        "(3) model performance metrics and a dollar-value simulation, "
        "(4) analysis of predicted-positive tickers, and "
        "(5) error analysis to understand model failure modes."
    )
    doc.add_page_break()

    # ── sections ──────────────────────────────────────────────────────────────
    print("Section 1: Trade Overview …")
    section_trade_overview(doc, df_full, proc_df)
    doc.add_page_break()

    print("Section 2: Feature EDA …")
    section_features_eda(doc, model, proc_df)
    doc.add_page_break()

    print("Section 3: Model Performance …")
    section_model_performance(
        doc, model, proc_df, test_df, y_test, y_test_cont, y_prob, xgb_clf
    )
    doc.add_page_break()

    print("Section 4: Ticker Analysis …")
    section_ticker_analysis(doc, test_df, y_test, y_test_cont, y_prob)
    doc.add_page_break()

    print("Section 5: Error Analysis …")
    section_error_analysis(doc, model, test_df, y_test, y_test_cont, y_prob)

    # ── save ──────────────────────────────────────────────────────────────────
    os.makedirs("data/output", exist_ok=True)
    doc.save(OUTPUT_PATH)
    size_mb = os.path.getsize(OUTPUT_PATH) / 1e6
    print(f"\nReport saved → {OUTPUT_PATH}  ({size_mb:.1f} MB)")


if __name__ == "__main__":
    main()
